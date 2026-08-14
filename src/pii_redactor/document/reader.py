import logging
from typing import List, Optional
from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from pii_redactor.models import DocumentBlock, RunInfo

logger = logging.getLogger(__name__)

class DocxReader:
    """
    Reads DOCX files and extracts text blocks (paragraphs, tables, headers, footers)
    while preserving run-level formatting metadata and character offsets.
    """

    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc: Optional[Document] = None

    def read(self) -> List[DocumentBlock]:
        """
        Parses the DOCX file and returns a list of DocumentBlocks.
        """
        try:
            self.doc = Document(self.file_path)
        except Exception as e:
            logger.error(f"Failed to open DOCX file {self.file_path}: {e}")
            raise ValueError(f"Invalid or corrupted DOCX file: {e}") from e

        blocks: List[DocumentBlock] = []

        # 1. Extract from Headers and Footers (all sections)
        for sec_idx, section in enumerate(self.doc.sections):
            # Headers
            if section.header is not None:
                header = section.header
                # Header paragraphs
                for p_idx, p in enumerate(header.paragraphs):
                    block_id = f"section_{sec_idx}_header_p_{p_idx}"
                    blocks.append(self._create_block(block_id, "header", p))
                # Header tables
                for t_idx, table in enumerate(header.tables):
                    blocks.extend(self._extract_from_table(table, f"section_{sec_idx}_header_table_{t_idx}"))

            # Footers
            if section.footer is not None:
                footer = section.footer
                # Footer paragraphs
                for p_idx, p in enumerate(footer.paragraphs):
                    block_id = f"section_{sec_idx}_footer_p_{p_idx}"
                    blocks.append(self._create_block(block_id, "footer", p))
                # Footer tables
                for t_idx, table in enumerate(footer.tables):
                    blocks.extend(self._extract_from_table(table, f"section_{sec_idx}_footer_table_{t_idx}"))

        # 2. Extract from Main Body Paragraphs
        for p_idx, p in enumerate(self.doc.paragraphs):
            block_id = f"body_p_{p_idx}"
            blocks.append(self._create_block(block_id, "paragraph", p))

        # 3. Extract from Main Body Tables
        for t_idx, table in enumerate(self.doc.tables):
            blocks.extend(self._extract_from_table(table, f"body_table_{t_idx}"))

        logger.info(f"Successfully read DOCX document. Extracted {len(blocks)} text blocks.")
        return blocks

    def _create_block(self, block_id: str, block_type: str, paragraph: Paragraph) -> DocumentBlock:
        """
        Assembles raw run text and generates character offset mappings.
        """
        unified_text = ""
        run_infos: List[RunInfo] = []

        for run in paragraph.runs:
            run_text = run.text
            if not run_text:
                continue
            start_offset = len(unified_text)
            end_offset = start_offset + len(run_text)
            unified_text += run_text
            run_infos.append(RunInfo(
                run=run,
                text=run_text,
                start=start_offset,
                end=end_offset
            ))

        # Fallback if python-docx reports empty runs but the paragraph has text
        if not unified_text and paragraph.text:
            unified_text = paragraph.text
            # Create a single run proxy if runs are missing
            if paragraph.runs:
                # If runs exist but text was empty, use them
                pass
            else:
                # If no runs exist, we cannot map runs. We'll add text in a new run.
                # Usually python-docx paragraphs have at least one run if there is text.
                pass

        return DocumentBlock(
            block_id=block_id,
            block_type=block_type,
            text=unified_text,
            element=paragraph,
            runs=run_infos
        )

    def _extract_from_table(self, table: Table, prefix: str) -> List[DocumentBlock]:
        """
        Recursively extracts paragraphs from cells in a table, including nested tables.
        """
        blocks: List[DocumentBlock] = []
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                # Extract paragraphs directly within the cell
                for p_idx, p in enumerate(cell.paragraphs):
                    block_id = f"{prefix}_r{row_idx}_c{col_idx}_p{p_idx}"
                    blocks.append(self._create_block(block_id, "table_cell", p))
                
                # Extract from nested tables inside this cell
                for nt_idx, nested_table in enumerate(cell.tables):
                    blocks.extend(self._extract_from_table(nested_table, f"{prefix}_r{row_idx}_c{col_idx}_nt_{nt_idx}"))
        return blocks
