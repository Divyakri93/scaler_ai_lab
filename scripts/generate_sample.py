import os
from docx import Document

def main():
    doc = Document()
    
    # 1. Setup Header & Footer
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = "CONFIDENTIAL: Client profile for Alice Smith (alice.smith@gmail.com)."
    
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = "Footer info - Page 1"

    # 2. Setup Body Paragraphs
    p0 = doc.add_paragraph()
    p0.add_run("Applicant Profile: ")
    p0.add_run("Alice Smith").bold = True
    p0.add_run(" (")
    p0.add_run("alice.smith@gmail.com")
    p0.add_run(") was born on ")
    p0.add_run("Date of Birth: 12/04/1999").italic = True
    p0.add_run(".")

    p1 = doc.add_paragraph()
    p1.add_run("Contact details: ")
    p1.add_run("+91 9876543210")
    p1.add_run(". IP: ")
    p1.add_run("192.168.1.1")
    p1.add_run(".")

    p2 = doc.add_paragraph()
    p2.add_run("SSN: ")
    p2.add_run("123-45-6789")
    p2.add_run(". Employed by ")
    p2.add_run("Google LLC")
    p2.add_run(".")

    p3 = doc.add_paragraph()
    p3.add_run("Address: ")
    p3.add_run("123 MG Road, Delhi")
    p3.add_run(". Payment card: ")
    p3.add_run("4111-1111-1111-1111")
    p3.add_run(".")

    # 3. Setup Table
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).paragraphs[0].text = "Alternative Address: Flat 201, Tower B, Sector 62, Noida"
    table.cell(0, 1).paragraphs[0].text = "Primary contact email: alice.smith@gmail.com"

    os.makedirs("input", exist_ok=True)
    doc.save("input/sample_document.docx")
    print("Successfully generated input/sample_document.docx")

if __name__ == "__main__":
    main()
