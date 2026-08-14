import os
import pytest
from docx import Document
from pii_redactor.config import RedactorConfig
from pii_redactor.document.processor import RedactionPipeline
from pii_redactor.document.reader import DocxReader
from pii_redactor.models import PIIType

def create_full_pii_docx(file_path: str):
    """
    Creates a sample DOCX containing all 9 PII categories.
    """
    doc = Document()
    
    # Header
    section = doc.sections[0]
    section.header.paragraphs[0].text = "CONFIDENTIAL: Client profile for Alice Smith (alice.smith@gmail.com)."
    
    # Body Paragraphs
    p1 = doc.add_paragraph()
    p1.add_run("The applicant is ")
    run_bold = p1.add_run("Alice Smith")
    run_bold.bold = True
    p1.add_run(". She was born on ")
    run_italic = p1.add_run("Date of Birth: 12/04/1999")
    run_italic.italic = True
    p1.add_run(". Contact: +91 9876543210 or visit local IP 192.168.1.1.")
    
    p2 = doc.add_paragraph("Her SSN is 123-45-6789. She works at Google LLC.")
    
    # Table containing remaining PII
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).paragraphs[0].text = "Address details: 123 MG Road, Delhi"
    table.cell(0, 1).paragraphs[0].text = "Credit Card details: 4111-1111-1111-1111"
    table.cell(1, 0).paragraphs[0].text = "Email backup: alice.smith@gmail.com"
    table.cell(1, 1).paragraphs[0].text = "Normal, non-PII text data."
    
    doc.save(file_path)

def test_full_redaction_pipeline(tmp_path):
    input_path = os.path.join(tmp_path, "input.docx")
    output_path = os.path.join(tmp_path, "output.docx")
    
    create_full_pii_docx(input_path)
    
    # Run the pipeline with seed 42
    config = RedactorConfig(seed=42)
    pipeline = RedactionPipeline(config=config)
    
    redacted_report = pipeline.redact(input_path, output_path)
    
    # 1. Verify the report generated entities
    # We should have found elements for all 9 categories
    flat_entities = []
    for entities in redacted_report.values():
        flat_entities.extend(entities)
        
    found_types = {e.entity_type for e in flat_entities}
    expected_types = {
        PIIType.FULL_NAME,
        PIIType.EMAIL,
        PIIType.DATE_OF_BIRTH,
        PIIType.PHONE,
        PIIType.IP_ADDRESS,
        PIIType.SSN,
        PIIType.COMPANY_NAME,
        PIIType.ADDRESS,
        PIIType.CREDIT_CARD
    }
    
    assert expected_types.issubset(found_types)
    
    # 2. Read redacted output and verify content
    output_reader = DocxReader(output_path)
    output_blocks = output_reader.read()
    
    # Check that original texts are completely gone and replaced
    for block in output_blocks:
        assert "Alice Smith" not in block.text
        assert "alice.smith@gmail.com" not in block.text
        assert "123-45-6789" not in block.text
        assert "4111-1111-1111-1111" not in block.text
        assert "192.168.1.1" not in block.text
        assert "12/04/1999" not in block.text
        assert "9876543210" not in block.text
        assert "123 MG Road, Delhi" not in block.text
        
        # Verify formatting was preserved on body paragraph
        if block.block_id.startswith("body_p_0"):
            # Find the bold run
            bold_runs = [r for r in block.runs if r.run.bold]
            assert len(bold_runs) > 0
            
            # Find the italic run
            italic_runs = [r for r in block.runs if r.run.italic]
            assert len(italic_runs) > 0
            
            # Verify consistent email replacement matches name replacement
            name_entity = next(e for e in flat_entities if e.entity_type == PIIType.FULL_NAME)
            email_entity = next(e for e in flat_entities if e.entity_type == PIIType.EMAIL)
            
            first_fake = name_entity.replacement.lower().split()[0]
            assert first_fake in email_entity.replacement.lower()
