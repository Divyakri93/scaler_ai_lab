import io
import pytest
from fastapi.testclient import TestClient
from docx import Document

from pii_redactor.api import app

client = TestClient(app)

def create_in_memory_docx() -> io.BytesIO:
    """
    Generates a mock docx in memory to test file upload endpoints.
    """
    doc = Document()
    doc.add_paragraph("Alice Smith email is alice@test.com.")
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def test_api_health_endpoint():
    """
    Verifies that the health check endpoint returns 200 and healthy status.
    """
    response = client.get("/health")
    assert response.status_code == 200
    assert response.json() == {"status": "healthy"}

def test_api_redact_endpoint():
    """
    Verifies uploading a docx, redacting it, and getting back the redacted file response.
    """
    file_stream = create_in_memory_docx()
    files = {"file": ("test_doc.docx", file_stream, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    
    # Run request
    response = client.post("/redact?seed=42", files=files)
    
    assert response.status_code == 200
    # Response content should be a valid docx file stream
    content = response.content
    assert len(content) > 0
    
    # Read response as docx and check that PII is redacted
    out_stream = io.BytesIO(content)
    doc_out = Document(out_stream)
    text = "".join(p.text for p in doc_out.paragraphs)
    
    assert "Alice Smith" not in text
    assert "alice@test.com" not in text

def test_api_redact_unsupported_file():
    """
    Verifies that uploading non-docx returns a 400 bad request.
    """
    files = {"file": ("test_doc.txt", b"some plain text data", "text/plain")}
    response = client.post("/redact", files=files)
    assert response.status_code == 400
    assert "Only .docx files are supported" in response.json()["detail"]
