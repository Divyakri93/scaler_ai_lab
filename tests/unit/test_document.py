import os
import pytest
from docx import Document
from pii_redactor.document.reader import DocxReader
from pii_redactor.document.writer import DocxWriter
from pii_redactor.document.processor import DocumentProcessor
from pii_redactor.models import PIIEntity, DocumentBlock

def create_sample_docx(file_path: str):
    """
    Helper to build a sample docx containing body, tables, header, footer.
    """
    doc = Document()
    
    # 1. Setup Header & Footer
    section = doc.sections[0]
    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = "Header text with John Doe email: john.doe@corp.com"
    
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = "Footer info - Page 1"

    # 2. Setup Body Paragraphs
    p1 = doc.add_paragraph()
    p1.add_run("My name is ")
    run_bold = p1.add_run("John Doe")
    run_bold.bold = True
    p1.add_run(" and I live here.")

    doc.add_paragraph("This is a plain paragraph without formatting.")

    # 3. Setup Table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).paragraphs[0].text = "Cell 0,0: John Doe"
    table.cell(0, 1).paragraphs[0].text = "Cell 0,1: normal text"
    
    # Nested table inside cell (1, 1)
    nested_table = table.cell(1, 1).add_table(rows=1, cols=1)
    nested_table.cell(0, 0).paragraphs[0].text = "Nested Cell: john.doe@corp.com"

    doc.save(file_path)

def test_document_reader_writer_processor(tmp_path):
    input_path = os.path.join(tmp_path, "input.docx")
    output_path = os.path.join(tmp_path, "output.docx")
    
    create_sample_docx(input_path)
    
    # Verify Reader
    reader = DocxReader(input_path)
    blocks = reader.read()
    
    # Verify we extracted all expected text blocks
    assert len(blocks) > 0
    
    # Find blocks containing "John Doe"
    john_blocks = [b for b in blocks if "John Doe" in b.text]
    assert len(john_blocks) >= 2  # One in body, one in table cell
    
    # Define a custom redactor callback that replaces "John Doe" with "Jane Smith"
    # and "john.doe@corp.com" with "jane.smith@corp.com"
    def mock_redact_callback(blocks_list):
        redacted = {}
        for block in blocks_list:
            entities = []
            if "John Doe" in block.text:
                idx = block.text.find("John Doe")
                entities.append(PIIEntity(
                    entity_type="FULL_NAME",
                    original_text="John Doe",
                    start=idx,
                    end=idx + len("John Doe"),
                    confidence=1.0,
                    source_detector="mock",
                    replacement="Jane Smith"
                ))
            if "john.doe@corp.com" in block.text:
                idx = block.text.find("john.doe@corp.com")
                entities.append(PIIEntity(
                    entity_type="EMAIL",
                    original_text="john.doe@corp.com",
                    start=idx,
                    end=idx + len("john.doe@corp.com"),
                    confidence=1.0,
                    source_detector="mock",
                    replacement="jane.smith@corp.com"
                ))
            if entities:
                redacted[block.block_id] = entities
        return redacted

    # Run Processor
    processor = DocumentProcessor()
    processor.process(input_path, output_path, mock_redact_callback)
    
    # Read redacted output and verify content
    output_reader = DocxReader(output_path)
    output_blocks = output_reader.read()
    
    # Check that "John Doe" has been replaced by "Jane Smith" and "john.doe@corp.com" replaced by "jane.smith@corp.com"
    for block in output_blocks:
        assert "John Doe" not in block.text
        assert "john.doe@corp.com" not in block.text
        
        # Verify substitutions
        if block.block_id.startswith("body_p_0"):
            assert "Jane Smith" in block.text
            # Verify bold formatting is preserved for "Jane Smith"
            runs_text = [r.run.text for r in block.runs]
            assert "Jane Smith" in runs_text
            
            # Find the run and check bold attribute
            jane_run_info = next(r for r in block.runs if "Jane Smith" in r.text)
            assert jane_run_info.run.bold is True
            
        if "header" in block.block_id:
            assert "jane.smith@corp.com" in block.text
